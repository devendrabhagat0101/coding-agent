"""
Build LLM context from a local project directory.

Produces two things:
  - An ASCII file-tree (always included)
  - Concatenated file contents, truncated at CONTEXT_MAX_CHARS

The combined result becomes the system-prompt preamble so the model
understands the project layout before any user message arrives.
"""

from __future__ import annotations

from pathlib import Path

from .config import CONTEXT_MAX_CHARS, IGNORED_DIRS, IGNORED_FILENAMES, TEXT_EXTENSIONS


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _is_ignored_dir(name: str) -> bool:
    return name in IGNORED_DIRS or name.endswith(".egg-info")


def _is_text_file(path: Path) -> bool:
    if path.name in IGNORED_FILENAMES:
        return False
    # .docx handled separately via _read_file
    if path.suffix.lower() == ".docx":
        return True
    # Extension match (handles extensionless names like Dockerfile too)
    return path.suffix in TEXT_EXTENSIONS or path.name in TEXT_EXTENSIONS


def _read_file(path: Path) -> str:
    """Read a file as text. Handles .docx by extracting paragraphs."""
    if path.suffix.lower() == ".docx":
        try:
            from docx import Document  # python-docx
            doc = Document(str(path))
            lines: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            return "\n".join(lines)
        except Exception:  # noqa: BLE001
            return f"[could not parse {path.name}]"
    return path.read_text(encoding="utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_file_tree(root: Path) -> str:
    """Return a pretty ASCII tree of *root*, skipping ignored paths."""
    lines: list[str] = [str(root.resolve())]

    def _walk(directory: Path, prefix: str) -> None:
        try:
            entries = sorted(
                directory.iterdir(),
                key=lambda p: (p.is_file(), p.name.lower()),
            )
        except PermissionError:
            return

        # Filter out ignored directories
        entries = [
            e for e in entries
            if not (e.is_dir() and _is_ignored_dir(e.name))
            and e.name not in IGNORED_FILENAMES
        ]

        for i, entry in enumerate(entries):
            connector = "└── " if i == len(entries) - 1 else "├── "
            lines.append(f"{prefix}{connector}{entry.name}")
            if entry.is_dir():
                extension = "    " if i == len(entries) - 1 else "│   "
                _walk(entry, prefix + extension)

    _walk(root, "")
    return "\n".join(lines)


def read_project_files(root: Path, max_chars: int = CONTEXT_MAX_CHARS) -> str:
    """
    Walk *root* and return the contents of all text files concatenated,
    capped at *max_chars* total characters.  Binary / ignored files are skipped.
    """
    parts: list[str] = []
    total = 0

    for path in sorted(root.rglob("*")):
        # Skip ignored directories anywhere in the path
        if any(_is_ignored_dir(part) for part in path.parts):
            continue
        if not path.is_file():
            continue
        if not _is_text_file(path):
            continue

        try:
            content = _read_file(path)
        except OSError:
            continue

        rel = path.relative_to(root)
        header = f"\n\n### {rel}\n"
        snippet = header + content

        if total + len(snippet) > max_chars:
            remaining = max_chars - total
            if remaining > len(header) + 80:
                parts.append(snippet[:remaining] + "\n... [truncated]")
            break

        parts.append(snippet)
        total += len(snippet)

    return "".join(parts)


def build_system_prompt(root: Path) -> str:
    """
    Combine file tree + file contents into a single system-prompt string
    that gives the LLM full project awareness.
    """
    tree = build_file_tree(root)
    files = read_project_files(root)

    return (
        "You are an expert coding assistant with full knowledge of the project below.\n"
        "Use this context to give precise, actionable answers.\n\n"
        "IMPORTANT — when asked to CREATE or WRITE files, use EXACTLY this format for each file:\n"
        "<!-- FILE: filename.ext -->\n"
        "```lang\n"
        "<full file content>\n"
        "```\n"
        "Replace 'lang' with the correct language tag (py, md, js, txt, etc.).\n"
        "This marker lets the agent write the file to disk automatically.\n\n"
        "## Project Structure\n\n"
        "```\n"
        f"{tree}\n"
        "```\n"
        "## File Contents\n"
        f"{files}\n"
    )
