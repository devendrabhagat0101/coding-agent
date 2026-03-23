"""
Autonomous Project Builder  (coding-agent v2.0.1)

Reads a requirements file (.md / .txt / .docx), asks Ollama to produce a
structured build plan, scaffolds each service, generates every file with
context-aware prompts, optionally reviews and auto-fixes each file, then
commits everything to git.

Usage (via CLI):
    coding-agent build requirements.md --dir ./my-project
    coding-agent build requirements.md --dir ./my-project --no-review --no-git
"""

from __future__ import annotations

import json
import os
import re
import subprocess
import textwrap
from pathlib import Path
from typing import Any

from rich.console import Console
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.syntax import Syntax
from rich.table import Table

from .engine import CodingEngine, OllamaUnavailableError
from . import git_ops

console = Console()
err_console = Console(stderr=True, style="bold red")

# ── Max chars of already-generated files to include as context ────────────────
GENERATED_CONTEXT_LIMIT = 20_000


# ═══════════════════════════════════════════════════════════════════════════════
# Requirements reader
# ═══════════════════════════════════════════════════════════════════════════════

def read_requirements(path: Path) -> str:
    """
    Read a requirements file.
    Supports: .md  .txt  .rst  .docx
    Returns plain text.
    """
    suffix = path.suffix.lower()

    if suffix == ".docx":
        try:
            from docx import Document  # python-docx
            doc = Document(str(path))
            lines: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text)
            # also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            return "\n".join(lines)
        except ImportError:
            raise RuntimeError(
                "python-docx is required to read .docx files.\n"
                "Install it: pip install python-docx"
            )
    else:
        return path.read_text(encoding="utf-8", errors="replace")


# ═══════════════════════════════════════════════════════════════════════════════
# Project Planner
# ═══════════════════════════════════════════════════════════════════════════════

_PLAN_SYSTEM = textwrap.dedent("""\
    You are a senior software architect. Given a requirements document you must
    produce a complete, structured JSON build plan.

    Rules:
    - Return ONLY a valid JSON object — no markdown fences, no explanations.
    - The JSON must exactly follow the schema below.
    - Keep file lists practical: only files that are truly needed.
    - scaffold_cmd must be a real shell command (or empty string "" if none needed).
    - For Flutter use: flutter create <name> --platforms web,android,ios
    - For Java Spring Boot use: "" (no scaffold; generate build.gradle + src tree manually)
    - For Python use: "" (no scaffold; generate files manually)

    ── Java Spring Boot file conventions ─────────────────────────────────────────
    Spring Boot 3.3.x · Java 21 · Gradle (build.gradle + settings.gradle)

    Standard directory structure (relative to the service output_dir):
      build.gradle
      settings.gradle
      src/main/java/{pkg}/                       (e.g. com/dev2/payment)
        {Name}Application.java                   @SpringBootApplication entry point
        config/
          OpenApiConfig.java                     springdoc OpenAPI bean
        controller/
          {Entity}Controller.java                @RestController, @RequestMapping("/api/{entity}")
        service/
          {Entity}Service.java                   @Service interface + Impl
        repository/
          {Entity}Repository.java                @Repository extends JpaRepository<{Entity}, Long>
        model/
          {Entity}.java                          @Entity, @Table, @Id, @GeneratedValue
        dto/
          {Entity}Request.java                   record or class with @Valid annotations
          {Entity}Response.java                  record mapping entity → response
          ApiResponse.java                       generic wrapper: success/error factory methods
        exception/
          ApiException.java                      extends RuntimeException + HttpStatus
          GlobalExceptionHandler.java            @RestControllerAdvice @ExceptionHandler
      src/main/resources/
        application.yml                          server.port, spring.datasource, logging
      src/test/java/{pkg}/
        {Name}ApplicationTests.java              @SpringBootTest contextLoads()
        controller/
          {Entity}ControllerTest.java            @WebMvcTest with @MockBean

    Key Spring Boot coding rules:
    - Use records for DTOs (Java 16+) — immutable, concise.
    - Return ResponseEntity<ApiResponse<T>> from all controllers.
    - Annotate service interfaces, not just impls.
    - Use @Valid on @RequestBody in controllers; validate in DTO with @NotNull/@Size.
    - Always include GlobalExceptionHandler to avoid raw 500s.
    - application.yml: set server.port, spring.application.name, datasource (if DB).
    - H2 in-memory for dev (spring.h2.console.enabled=true).
    - build.gradle dependencies: spring-boot-starter-web, actuator, validation,
      data-jpa (if DB), h2/postgresql/mysql (runtime), springdoc-openapi (2.6.0).
    - springdoc version: 2.6.0 (springdoc-openapi-starter-webmvc-ui).
    - Spring Boot version: 3.3.5 · dependency-management: 1.1.6.

    JSON schema:
    {
      "project_name": "string",
      "description": "string",
      "services": [
        {
          "name": "string",
          "tech": "Flutter | Java Spring Boot | Python FastAPI | Python | Other",
          "language": "dart | java | python | other",
          "output_dir": "string  (relative folder name, e.g. weather_ui)",
          "scaffold_cmd": "string  (shell command or empty string)",
          "dependencies": ["string"],
          "files": [
            {
              "path": "string  (relative to service output_dir)",
              "description": "string  (what this file does)",
              "prompt": "string  (detailed instruction for code generation)"
            }
          ]
        }
      ]
    }
""")


def _extract_json(text: str) -> str:
    """Strip markdown fences and extract the first JSON object."""
    # Remove ```json ... ``` fences
    text = re.sub(r"```(?:json)?\s*", "", text)
    text = re.sub(r"```", "", text)
    # Find first { ... } block
    start = text.find("{")
    if start == -1:
        return text.strip()
    depth = 0
    for i, ch in enumerate(text[start:], start):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start : i + 1]
    return text[start:].strip()


def plan_project(requirements: str, engine: CodingEngine) -> dict[str, Any]:
    """
    Ask Ollama to produce a JSON build plan from the requirements text.
    Returns the parsed dict.
    """
    prompt = (
        "Produce the build plan JSON for the following requirements.\n\n"
        "REQUIREMENTS:\n"
        "─────────────\n"
        f"{requirements[:6000]}\n"   # cap to avoid overflowing context
        "─────────────\n\n"
        "Return ONLY the JSON object, nothing else."
    )

    with Progress(SpinnerColumn(), TextColumn("[cyan]{task.description}"), transient=True) as p:
        p.add_task("AI is planning the project structure …")
        raw = engine.complete(prompt, system=_PLAN_SYSTEM, temperature=0.1)

    json_str = _extract_json(raw)
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as exc:
        raise RuntimeError(
            f"AI returned invalid JSON for the project plan.\n"
            f"Raw output:\n{raw[:800]}\n\nError: {exc}"
        ) from exc


# ═══════════════════════════════════════════════════════════════════════════════
# Scaffolding
# ═══════════════════════════════════════════════════════════════════════════════

def scaffold_service(cmd: str, cwd: Path) -> bool:
    """
    Run a scaffold shell command (e.g. `flutter create weather_ui`).
    Returns True on success.
    """
    if not cmd:
        cwd.mkdir(parents=True, exist_ok=True)
        return True

    console.print(f"\n[bold yellow]Scaffold:[/] [cyan]{cmd}[/]")
    console.print(f"[dim]  in: {cwd}[/]")

    cwd.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        cmd,
        shell=True,
        cwd=str(cwd),
        capture_output=False,
        text=True,
    )
    if result.returncode != 0:
        err_console.print(f"✗ Scaffold command failed (exit {result.returncode})")
        return False
    return True


# ═══════════════════════════════════════════════════════════════════════════════
# File Generator
# ═══════════════════════════════════════════════════════════════════════════════

_FILE_SYSTEM = textwrap.dedent("""\
    You are an expert software engineer.
    Your job is to generate complete, production-quality source code for a single file.

    Rules:
    - Return ONLY the file content inside a single fenced code block.
    - Do NOT add any explanation outside the code block.
    - The code must be complete — no placeholders like "// TODO" or "pass".
    - Follow best practices for the language/framework.
    - Match the existing project structure and naming conventions shown in context.
""")


def generate_file(
    file_info: dict[str, str],
    service: dict[str, Any],
    requirements_summary: str,
    generated_so_far: dict[str, str],
    engine: CodingEngine,
) -> str:
    """
    Generate code for a single file.
    Returns the raw file content (no fences).
    """
    path = file_info["path"]
    description = file_info["description"]
    instruction = file_info.get("prompt", description)
    lang = Path(path).suffix.lstrip(".") or "text"

    # Build context from already-generated files (capped)
    ctx_parts: list[str] = []
    total = 0
    for p, content in generated_so_far.items():
        snippet = f"\n### {p}\n```\n{content}\n```\n"
        if total + len(snippet) > GENERATED_CONTEXT_LIMIT:
            break
        ctx_parts.append(snippet)
        total += len(snippet)
    context_block = "".join(ctx_parts)

    prompt = (
        f"Project: {service.get('name', '')} ({service.get('tech', '')})\n\n"
        f"Requirements summary:\n{requirements_summary[:2000]}\n\n"
        f"File to generate: {path}\n"
        f"Purpose: {description}\n\n"
        f"Instruction: {instruction}\n\n"
        + (f"Already generated files for context:\n{context_block}\n\n" if context_block else "")
        + f"Generate the complete content of `{path}` ({lang}).\n"
        "Return ONLY the file content inside a single fenced code block."
    )

    with Progress(SpinnerColumn(), TextColumn(f"[cyan]Generating {path} …"), transient=True) as p_prog:
        p_prog.add_task("")
        raw = engine.complete(prompt, system=_FILE_SYSTEM, temperature=0.15)

    # Extract code from fenced block
    blocks = re.findall(r"```(?:\w*)\n(.*?)```", raw, re.DOTALL)
    if blocks:
        return blocks[0].rstrip() + "\n"
    # Fallback: return raw (model may have skipped fences)
    return raw.strip() + "\n"


# ═══════════════════════════════════════════════════════════════════════════════
# Reviewer / Auto-fixer
# ═══════════════════════════════════════════════════════════════════════════════

_REVIEW_SYSTEM = textwrap.dedent("""\
    You are a senior code reviewer.
    Given a source file, identify real bugs, missing imports, obvious errors,
    or incomplete implementations.
    If the file is correct, reply with exactly: NO_ISSUES
    Otherwise, return the COMPLETE corrected file inside a single fenced code block.
    Do not add explanations outside the block.
""")


def review_and_fix(path: str, content: str, engine: CodingEngine) -> str:
    """
    Review a generated file and return fixed content.
    Returns original content if no issues found.
    """
    lang = Path(path).suffix.lstrip(".") or "text"
    prompt = (
        f"Review and fix `{path}` if needed.\n\n"
        f"```{lang}\n{content}\n```\n\n"
        "If correct: reply NO_ISSUES\n"
        "If issues found: return the complete fixed file in a fenced code block."
    )

    with Progress(SpinnerColumn(), TextColumn(f"[dim]Reviewing {path} …"), transient=True) as p:
        p.add_task("")
        reply = engine.complete(prompt, system=_REVIEW_SYSTEM, temperature=0.1)

    if "NO_ISSUES" in reply.upper():
        return content

    blocks = re.findall(r"```(?:\w*)\n(.*?)```", reply, re.DOTALL)
    if blocks:
        return blocks[0].rstrip() + "\n"
    return content


# ═══════════════════════════════════════════════════════════════════════════════
# Main Orchestrator
# ═══════════════════════════════════════════════════════════════════════════════

def build(
    req_file: Path,
    output_dir: Path,
    model: str,
    do_review: bool = True,
    do_git: bool = True,
) -> None:
    """
    Full autonomous build:
      1. Read requirements
      2. Plan project (AI → JSON)
      3. For each service: scaffold → generate files → review+fix → write
      4. Git init + commit
    """
    # ── 1. Read requirements ──────────────────────────────────────────────────
    console.print(f"\n[bold cyan]Reading requirements:[/] {req_file}")
    requirements = read_requirements(req_file)
    requirements_summary = requirements[:3000]
    console.print(f"[dim]{len(requirements)} chars loaded.[/]")

    # ── 2. Engine check ───────────────────────────────────────────────────────
    engine = CodingEngine(model=model)
    available = engine.list_local_models()
    if not available:
        raise OllamaUnavailableError(
            "Cannot reach Ollama. Run: ollama serve"
        )
    if model not in available:
        raise OllamaUnavailableError(
            f"Model '{model}' not found. Available: {', '.join(available)}\n"
            f"Pull it: ollama pull {model}"
        )

    # ── 3. Plan ───────────────────────────────────────────────────────────────
    plan = plan_project(requirements, engine)

    project_name = plan.get("project_name", "project")
    services     = plan.get("services", [])

    _print_plan(plan)

    output_dir.mkdir(parents=True, exist_ok=True)

    total_files  = sum(len(s.get("files", [])) for s in services)
    written      = 0
    all_paths: list[Path] = []

    # ── 4. Build each service ─────────────────────────────────────────────────
    for svc in services:
        svc_name    = svc.get("name", "service")
        svc_tech    = svc.get("tech", "")
        svc_outdir  = svc.get("output_dir", svc_name)
        scaffold_cmd = svc.get("scaffold_cmd", "")
        files        = svc.get("files", [])

        svc_root = output_dir / svc_outdir

        console.print(
            Panel(
                f"[bold]{svc_name}[/]  •  [cyan]{svc_tech}[/]  •  {len(files)} files\n"
                f"[dim]→ {svc_root}[/]",
                title="[bold green]Building service[/]",
                expand=False,
            )
        )

        # Scaffold
        scaffold_service(scaffold_cmd, svc_root)

        # Generate files
        generated: dict[str, str] = {}   # path → content (for cross-file context)

        for file_info in files:
            rel_path = file_info.get("path", "")
            if not rel_path:
                continue

            abs_path = svc_root / rel_path

            # Generate
            content = generate_file(
                file_info=file_info,
                service=svc,
                requirements_summary=requirements_summary,
                generated_so_far=generated,
                engine=engine,
            )

            # Review + fix
            if do_review:
                fixed = review_and_fix(rel_path, content, engine)
                if fixed != content:
                    console.print(f"  [yellow]↻[/] Fixed issues in [cyan]{rel_path}[/]")
                    content = fixed
                else:
                    console.print(f"  [green]✓[/] [cyan]{rel_path}[/]")
            else:
                console.print(f"  [green]✓[/] [cyan]{rel_path}[/]")

            # Write
            abs_path.parent.mkdir(parents=True, exist_ok=True)
            abs_path.write_text(content, encoding="utf-8")
            all_paths.append(abs_path)
            generated[rel_path] = content
            written += 1

        console.print(f"[dim]  {len(files)} files written to {svc_root}[/]")

    # ── 5. Git ────────────────────────────────────────────────────────────────
    if do_git and all_paths:
        console.print("\n[bold cyan]Initialising git repository …[/]")
        try:
            repo = git_ops.get_repo(output_dir)
            if repo is None:
                git_ops.init_repo(output_dir)

            # Stage all written files
            import git as _git
            repo2 = _git.Repo(str(output_dir), search_parent_directories=True)
            repo2.git.add(A=True)

            if repo2.index.diff("HEAD") or repo2.untracked_files:
                repo2.index.commit(
                    f"feat: initial {project_name} scaffold generated by coding-agent v2.0.1"
                )
                console.print("[green]✓[/] Committed all generated files.")
        except Exception as exc:  # noqa: BLE001
            console.print(f"[yellow]⚠ Git commit skipped: {exc}[/]")

    # ── 6. Summary ────────────────────────────────────────────────────────────
    _print_summary(plan, output_dir, written, total_files)


# ═══════════════════════════════════════════════════════════════════════════════
# Rich display helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _print_plan(plan: dict[str, Any]) -> None:
    console.print(
        Panel(
            f"[bold]{plan.get('project_name', '?')}[/]\n"
            f"[dim]{plan.get('description', '')}[/]",
            title="[bold green]Build Plan[/]",
            expand=False,
        )
    )
    table = Table(show_header=True, header_style="bold cyan")
    table.add_column("Service")
    table.add_column("Tech")
    table.add_column("Output Dir")
    table.add_column("Files", justify="right")

    for svc in plan.get("services", []):
        table.add_row(
            svc.get("name", ""),
            svc.get("tech", ""),
            svc.get("output_dir", ""),
            str(len(svc.get("files", []))),
        )
    console.print(table)
    console.print()


def _print_summary(plan: dict, output_dir: Path, written: int, total: int) -> None:
    console.print(
        Panel(
            f"[bold green]Build complete![/]\n\n"
            f"  Project  : [cyan]{plan.get('project_name', '?')}[/]\n"
            f"  Output   : [cyan]{output_dir.resolve()}[/]\n"
            f"  Files    : [cyan]{written}/{total}[/] generated\n\n"
            "[dim]Next steps:[/]\n"
            "  1. Review generated files\n"
            "  2. Add a git remote:  /remote add origin <url>\n"
            "  3. Push:              /push",
            title="[bold green]coding-agent v2.0.1[/]",
            expand=False,
        )
    )
